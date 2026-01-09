from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path
import pypandoc

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None

try:
    from IPython.display import Image, display
except ImportError:  # pragma: no cover - optional dependency for notebooks
    Image = None
    display = None

from typing import Any, Dict, List, Tuple

from langchain_openai import ChatOpenAI
# from langchain_anthropic import ChatAnthropic
from langchain_google_genai import ChatGoogleGenerativeAI

from langgraph.prebuilt import ToolNode
from langgraph.graph import add_messages

from dotenv import load_dotenv

from langchain_core.messages import messages_from_dict, messages_to_dict

from apps.api.app.services.process_design_agents.default_config import DEFAULT_CONFIG
from apps.api.app.services.process_design_agents.agents.utils.agent_sizing_tools import (
    size_heat_exchanger_basic,
    size_pump_basic,
)
from apps.api.app.services.process_design_agents.agents.utils.equipment_stream_markdown import (
    equipments_and_streams_dict_to_markdown,
)
from apps.api.app.services.process_design_agents.utils.pydantic_utils import (
    EquipmentAndStreamList,
)

from .setup import GraphSetup
from .propagator import Propagator
from langgraph.checkpoint.memory import MemorySaver

load_dotenv()

class ProcessDesignGraph:
    """Main class that orchestractes the process design workflow."""
    
    def __init__(
        self,
        debug: bool = False,  # debug mode use with cli
        config: Dict[str, Any] = None,  # config dictionary
        delay_time: float = 0.5,  # set up the delay time in second
        save_graph_image: bool = False,
        graph_image_filename: str = "graph.png",
    ):
        """Initialize the process design agents graph and component.
        Args:
            debug: Whether to run in debug mode
            config: Configuration dictionary
        """
        self.debug = debug
        self.config = config or DEFAULT_CONFIG
        
        # a response_format for equipment and stream list output from llm
        self.response_format = {}
        self.save_graph_image = save_graph_image
        self.graph_image_filename = graph_image_filename
        
        # Initialize LLMs
        self.deep_thinking_llm = None
        self.quick_thinking_llm = None
        self.deep_structured_llm = None
        self.quick_structured_llm = None
        
        # Initialize LLMs by LLM provider.
        # if self.config["llm_provider"].lower() == "openai" or self.config["llm_provider"] == "ollama" or self.config["llm_provider"] == "openrouter":
        if self.config["llm_provider"].lower() == "openai":
            base_url = self._get_url_by_name(self.config["llm_provider"].lower())
            api_key = os.getenv("OPENAI_API_KEY")
            
            # Get the JSON schema from the Pydanitc model
            schema = EquipmentAndStreamList.model_json_schema()
            
            # Todo: Create llms with non-structured and structured output
            # model_kwargs={
            #     "text_format" = EquipmentAndStreamList 
            # }
            
        elif self.config["llm_provider"].lower() == "openrouter":
            base_url = self._get_url_by_name(self.config["llm_provider"].lower())
            api_key = os.getenv("OPENROUTER_API_KEY")
            
            # Get the JSON schema from the Pydanitc model
            schema = EquipmentAndStreamList.model_json_schema()
            
            # Build the exact response_format object OpenRouter expects
            self.response_format = {
                "type": "json_schema",
                "json_schema": {
                    "name": "equipment_and_stream_results",
                    "strict": True,
                    "schema": schema
                }
            }
            
            # Initialize ChatOpenAI to use OpenRouter
            self.deep_thinking_llm = ChatOpenAI(
                base_url=base_url,
                api_key=api_key,
                model=self.config["deep_think_llm"],
            )
            self.quick_thinking_llm = ChatOpenAI(
                base_url=base_url,
                api_key=api_key,
                model=self.config["quick_think_llm"],
            )
            self.deep_structured_llm = ChatOpenAI(
                base_url=base_url,
                api_key=api_key,
                model=self.config["deep_think_llm"],
                model_kwargs={
                    "response_format": self.response_format
                    }
            )
            self.quick_structured_llm = ChatOpenAI(
                base_url=base_url,
                api_key=api_key,
                model=self.config["quick_think_llm"],
                model_kwargs={
                    "response_format": self.response_format
                    }
            )
        elif self.config["llm_provider"].lower() == "ollama":
            base_url = self._get_url_by_name(self.config["llm_provider"].lower())
            
            # Get the JSON schema from the Pydanitc model
            schema = EquipmentAndStreamList.model_json_schema()
            
            self.deep_thinking_llm = ChatOpenAI(
                base_url=base_url,
                model=self.config["deep_think_llm"],
            )
            self.quick_thinking_llm = ChatOpenAI(
                base_url=base_url,
                model=self.config["quick_think_llm"]
            )
            self.deep_structured_llm = ChatOpenAI(
                base_url=base_url,
                model=self.config["deep_think_llm"],
                model_kwargs={
                    "format": schema,
                }
            )
            self.quick_structured_llm = ChatOpenAI(
                base_url=base_url,
                model=self.config["quick_think_llm"],
                model_kwargs={
                    "format": schema,
                }
            )
        # elif self.config["llm_provider"].lower() == "anthropic":
        #     self.deep_thinking_llm = ChatAnthropic(model=self.config["deep_think_llm"], base_url=self.config["backend_url"])
        #     self.quick_thinking_llm = ChatAnthropic(model=self.config["quick_think_llm"], base_url=self.config["backend_url"])
        elif self.config["llm_provider"].lower() == "google":
            
            # 
            self.response_format = {
                "config": {
                    "response_mime_type": "application/json",
                    "response_schema": EquipmentAndStreamList,       
                }
            }
            
            self.deep_thinking_llm = ChatGoogleGenerativeAI(model=self.config["deep_think_llm"])
            self.quick_thinking_llm = ChatGoogleGenerativeAI(model=self.config["quick_think_llm"])
            self.deep_structured_llm = ChatGoogleGenerativeAI(model=self.config["deep_think_llm"], model_kwargs=self.response_format)
            self.quick_structured_llm = ChatGoogleGenerativeAI(model=self.config["quick_think_llm"], model_kwargs=self.response_format)
        else:
            raise ValueError(f"Unsupported LLM provider: {self.config['llm_provider']}")

        # Set temperature for LLM
        self.deep_thinking_llm.temperature = self.config["deep_think_temperature"]
        self.quick_thinking_llm.temperature = self.config["quick_think_temperature"]
        self.deep_structured_llm.temperature = self.config["deep_think_temperature"]
        self.quick_structured_llm.temperature = self.config["quick_think_temperature"]
        
        # Initialize checkpointer
        self.checkpointer = MemorySaver()
        
        # Create tool nodes
        self.tool_nodes = self._create_tool_nodes()

        # Create the graph
        self.graph_setup = GraphSetup(
            llm_provider=self.config["llm_provider"].lower(),
            quick_thinking_llm=self.quick_thinking_llm,
            deep_thinking_llm=self.deep_thinking_llm,
            quick_structured_llm=self.quick_structured_llm,
            deep_structured_llm=self.deep_structured_llm,
            tool_nodes=self.tool_nodes,
            checkpointer=self.checkpointer,
            delay_time=delay_time,
        )
        
        # Initialize the propagator
        self.propagator = Propagator()
        
        # Initialize state
        self.problem_statement = None
        self.log_state_dict = {}
        
        # Set up the graph
        self.graph = self.graph_setup.setup_graph()
        self.agent_execution_order: List[Tuple[str, Any]] = self.graph_setup.get_agent_execution_order()
        
        # Print the graph
        self._render_graph_image()
        
        # Set log paths
        self.current_state_log_path = Path("eval_results/ProcessDesignAgents_logs/current_state_log.json")
        
    def propagate(
        self,
        problem_statement: str = "",
        save_markdown: str | None = None,
        save_word_doc: str | None = None,
        manual_concept_selection: bool = False,
        resume_from_last_run: bool = True,
    ):
        """Run the agent graph for a problem statement.

        Args:
            problem_statement: Design brief to analyse.
            save_markdown: Optional path for saving the aggregated markdown report.
            manual_concept_selection: When True, prompt the user to choose a concept
                instead of automatically selecting the highest feasibility score.
            resume_from_last_run: When True (default), continue from the last incomplete run
                for the same problem statement if a checkpoint exists.
        """
        
        # Set the prompt statement from user
        self.problem_statement = problem_statement
        
        # Set the concept selection provider
        previous_provider = self.graph_setup.concept_selection_provider
        
        # If manual selection is enable, then create selecting function.
        if manual_concept_selection:
            def _prompt_user(concept_options):
                print("\nSelect a concept for detailed development:", flush=True)
                for index, option in enumerate(concept_options, start=1):
                    score = option["score"]
                    score_text = f"{score}" if score is not None else "N/A"
                    print(f"{index}. {option['title']} (Feasibility Score: {score_text})", flush=True)
                choice = input("Enter choice (press Enter to use highest score): ").strip()
                if not choice:
                    return None
                try:
                    selection = int(choice)
                except ValueError:
                    print("Invalid selection. Defaulting to highest score.", flush=True)
                    return None
                if 1 <= selection <= len(concept_options):
                    return selection - 1
                print("Selection out of range. Defaulting to highest score.", flush=True)
                return None

            self.graph_setup.concept_selection_provider = _prompt_user
        else:
            # self.graph_setup.concept_selection_provider = None
            def _auto_provider(concept_options):
                return None

            self.graph_setup.concept_selection_provider = _auto_provider


        current_state, completed_agents, agent_outputs, resume_enabled = self._prepare_initial_state(
            problem_statement, resume_from_last_run
        )
        completed_agents = list(completed_agents)
        agent_outputs = dict(agent_outputs)
        pending_agents = [name for name, _ in self.agent_execution_order if name not in completed_agents]

        is_complete = False
        try:
            print(f"\n=========================== Start Line ===========================", flush=True)
            print(f"LLM Provider: {self.config['llm_provider']}", flush=True)
            print(f"Quick Thinking LLM: {self.config['quick_think_llm']}", flush=True)
            print(f"Deep Thinking LLM: {self.config['deep_think_llm']}", flush=True)
            if resume_enabled and completed_agents:
                if pending_agents:
                    print(
                        f"Resuming from agent: {pending_agents[0]} "
                        f"({len(completed_agents)} completed)",
                        flush=True,
                    )
                else:
                    print("No pending agents detected; validating stored results.", flush=True)
            print(f"=================================================================\n", flush=True)

            for agent_name, agent_fn in self.agent_execution_order:
                if agent_name in completed_agents:
                    continue

                try:
                    agent_result = agent_fn(current_state) or {}
                    if not isinstance(agent_result, dict):
                        agent_result = {}
                except Exception:
                    # Persist current progress before propagating the exception
                    self._save_current_state_log(
                        problem_statement=problem_statement,
                        current_state=current_state,
                        completed_agents=completed_agents,
                        agent_outputs=agent_outputs,
                        is_complete=False,
                    )
                    raise

                self._merge_state_updates(current_state, agent_result)
                serialized_output = (
                    self._serialize_state_dict(agent_result) if agent_result else {}
                )
                agent_outputs[agent_name] = serialized_output
                completed_agents.append(agent_name)

                self._save_current_state_log(
                    problem_statement=problem_statement,
                    current_state=current_state,
                    completed_agents=completed_agents,
                    agent_outputs=agent_outputs,
                    is_complete=False,
                )

            is_complete = True
            print(f"\n=========================== Finish Line ===========================", flush=True)
        finally:
            self.graph_setup.concept_selection_provider = previous_provider
            self._save_current_state_log(
                problem_statement=problem_statement,
                current_state=current_state,
                completed_agents=completed_agents,
                agent_outputs=agent_outputs,
                is_complete=is_complete,
            )
        
        # Store current state for reflection
        self.curr_state = current_state
        
        # Log state to default location
        self._log_state(current_state)

        if save_markdown:
            self._write_markdown_report(current_state, save_markdown)
            
        if save_word_doc:
            self._write_word_report(current_state, save_word_doc)
        
        return current_state
        
    def _create_tool_nodes(self) -> Dict[str, ToolNode]:
        """Create tool nodes for different equipment using abstract methods."""
        return {
            "equipment_sizing": ToolNode(
                [
                    size_heat_exchanger_basic,
                    size_pump_basic
                ]
            ),
        }
    
    def _render_graph_image(self) -> None:
        """Optionally render and persist the compiled graph to a PNG image."""
        if not self.save_graph_image:
            return

        app = getattr(self, "graph", None)
        if app is None:
            return

        try:
            self.graph.get_graph().print_ascii()
            png_bytes = app.get_graph().draw_png()
        except Exception as exc:  # pragma: no cover - visualization helper
            if self.debug:
                print(f"Unable to render graph diagram: {exc}", flush=True)
            return

        output_path = Path(self.graph_image_filename)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        try:
            output_path.write_bytes(png_bytes)
            if self.debug:
                print(f"Graph PNG saved to: {output_path.resolve()}", flush=True)
        except OSError as exc:  # pragma: no cover - filesystem error
            if self.debug:
                print(f"Failed to write graph PNG: {exc}", flush=True)

        if Image is not None and display is not None:
            try:
                display(Image(data=png_bytes))
            except Exception as exc:  # pragma: no cover - display helper
                if self.debug:
                    print(f"Unable to display graph image: {exc}", flush=True)

    # ------------------------------------------------------------------ #
    # Logging and resume helpers
    # ------------------------------------------------------------------ #

    def _make_json_safe(self, value: Any) -> Any:
        """Ensure values are JSON serializable."""
        if isinstance(value, dict):
            return {k: self._make_json_safe(v) for k, v in value.items()}
        if isinstance(value, list):
            return [self._make_json_safe(v) for v in value]
        if isinstance(value, (str, int, float, bool)) or value is None:
            return value
        return str(value)

    def _serialize_state_dict(self, state: Dict[str, Any]) -> Dict[str, Any]:
        """Convert a design state into a JSON-safe dictionary."""
        serialized: Dict[str, Any] = {}
        for key, value in state.items():
            if key == "messages" and isinstance(value, list):
                serialized[key] = messages_to_dict(value)
            else:
                serialized[key] = self._make_json_safe(value)
        return serialized

    def _deserialize_state_dict(self, state_data: Dict[str, Any]) -> Dict[str, Any]:
        """Restore a design state dictionary from JSON-safe payload."""
        restored: Dict[str, Any] = dict(state_data or {})
        messages_payload = restored.get("messages")
        if messages_payload is not None:
            try:
                restored["messages"] = messages_from_dict(messages_payload)
            except Exception:
                restored["messages"] = []
        else:
            restored["messages"] = restored.get("messages", [])
        return restored

    def _merge_state_updates(self, state: Dict[str, Any], updates: Dict[str, Any]) -> Dict[str, Any]:
        """Apply partial updates from an agent into the current state."""
        if not updates:
            return state
        for key, value in updates.items():
            if key == "messages" and value is not None:
                existing_messages = state.get("messages", [])
                if not isinstance(existing_messages, list):
                    existing_messages = list(existing_messages) if existing_messages else []
                state["messages"] = add_messages(existing_messages, value)
            else:
                state[key] = value
        return state

    def _load_current_state_log(self) -> Dict[str, Any] | None:
        """Load the on-disk resume log if it exists."""
        if not self.current_state_log_path.exists():
            return None
        try:
            return json.loads(self.current_state_log_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            return None

    def _save_current_state_log(
        self,
        *,
        problem_statement: str,
        current_state: Dict[str, Any],
        completed_agents: List[str],
        agent_outputs: Dict[str, Dict[str, Any]],
        is_complete: bool,
    ) -> None:
        """Persist the incremental agent progress for resume capability."""
        serialized_state = self._serialize_state_dict(current_state)
        log_payload = {
            "problem_statement": problem_statement,
            "agent_order": [name for name, _ in self.agent_execution_order],
            "completed_agents": completed_agents,
            "pending_agents": [
                name for name, _ in self.agent_execution_order if name not in completed_agents
            ],
            "last_completed_agent": completed_agents[-1] if completed_agents else None,
            "current_state": serialized_state,
            "agent_outputs": {
                name: self._make_json_safe(value) for name, value in agent_outputs.items()
            },
            "is_complete": is_complete,
            "updated_at": datetime.utcnow().isoformat() + "Z",
        }
        self.current_state_log_path.parent.mkdir(parents=True, exist_ok=True)
        self.current_state_log_path.write_text(
            json.dumps(log_payload, indent=4),
            encoding="utf-8",
        )

    def _prepare_initial_state(
        self, problem_statement: str, resume_from_last_run: bool
    ) -> tuple[Dict[str, Any], List[str], Dict[str, Dict[str, Any]], bool]:
        """Load existing progress if available, otherwise create a fresh state."""
        log_data = self._load_current_state_log() if resume_from_last_run else None
        resume_enabled = False
        completed_agents: List[str] = []
        agent_outputs: Dict[str, Dict[str, Any]] = {}

        if (
            log_data
            and log_data.get("problem_statement") == problem_statement
            and not log_data.get("is_complete", False)
        ):
            resume_enabled = True
            current_state = self._deserialize_state_dict(
                log_data.get("current_state", {})
            )
            completed_agents = list(dict.fromkeys(log_data.get("completed_agents", [])))
            valid_agent_names = {name for name, _ in self.agent_execution_order}
            completed_agents = [name for name in completed_agents if name in valid_agent_names]
            agent_outputs = {
                name: value
                for name, value in dict(log_data.get("agent_outputs", {})).items()
                if name in valid_agent_names
            }
        else:
            current_state = self.propagator.create_initial_state(problem_statement)
            current_state["llm_provider"] = self.config["llm_provider"]
            resume_enabled = False
            completed_agents = []
            agent_outputs = {}

        current_state.setdefault("llm_provider", self.config["llm_provider"])
        current_state.setdefault("messages", current_state.get("messages", []))

        self._save_current_state_log(
            problem_statement=problem_statement,
            current_state=current_state,
            completed_agents=completed_agents,
            agent_outputs=agent_outputs,
            is_complete=False,
        )

        return current_state, completed_agents, agent_outputs, resume_enabled
        
    def _get_url_by_name(self, name: str) -> str:
        """
        Retrieves the URL associated with a given name from a predefined list of base URLs.

        Args:
            name (str): The name of the API provider (e.g., 'OpenAI', 'Anthropic').

        Returns:
            str | None: The corresponding URL if found, otherwise None.
        """
        BASE_URLS = [
            ("openai", "https://api.openai.com/v1"),
            ("anthropic", "https://api.anthropic.com/"),
            ("google", "https://generativelanguage.googleapis.com/v1"),
            ("openrouter", "https://openrouter.ai/api/v1"),
            ("ollama", "http://localhost:11434/v1"),        
        ]
        
        for provider, url in BASE_URLS:
            if provider.lower() == name.lower():
                return url
        return None
    
    def _log_state(self, final_state):
        """Log the final state to a JSON file."""
        self.log_state_dict = {
            "problem_statement": self._make_json_safe(final_state.get("problem_statement", "")),
            "process_requirements": self._make_json_safe(final_state.get("process_requirements", "")),
            "research_concepts": self._make_json_safe(final_state.get("research_concepts", "")),
            "selected_concept_name": self._make_json_safe(final_state.get("selected_concept_name", "")),
            "selected_concept_details": self._make_json_safe(final_state.get("selected_concept_details", "")),
            "design_basis": self._make_json_safe(final_state.get("design_basis", "")),
            "flowsheet_description": self._make_json_safe(final_state.get("flowsheet_description", "")),
            "stream_list_template": self._make_json_safe(final_state.get("stream_list_template", "")),
            "stream_list_results": self._make_json_safe(final_state.get("stream_list_results", "")),
            "equipment_list_template": self._make_json_safe(final_state.get("equipment_list_template", "")),
            "equipment_list_results": self._make_json_safe(final_state.get("equipment_list_results", "")),
            "equipment_and_stream_template": self._make_json_safe(final_state.get("equipment_and_stream_template", "")),
            "equipment_and_stream_results": self._make_json_safe(final_state.get("equipment_and_stream_results", "")),
            "safety_risk_analyst_report": self._make_json_safe(final_state.get("safety_risk_analyst_report", "")),
            "project_manager_report": self._make_json_safe(final_state.get("project_manager_report", "")),
            "project_approval": self._make_json_safe(final_state.get("project_approval", "")),
        }
        
        # Save to file
        directory = Path(f"eval_results/ProcessDesignAgents_logs/")
        directory.mkdir(parents=True, exist_ok=True)
        
        with open(
            f"eval_results/ProcessDesignAgents_logs/full_states_log.json", "w"
        ) as f:
            json.dump(self.log_state_dict, f, indent=4)
        
    def _compose_report_sections(self, final_state: Dict[str, Any]) -> list[tuple[str, str]]:
        raw_equipment_and_streams = final_state.get("equipment_and_stream_results", "")
        if isinstance(raw_equipment_and_streams, str):
            raw_equipment_and_streams = json.loads(raw_equipment_and_streams)
            equipment_and_streams_markdown, _, _ = equipments_and_streams_dict_to_markdown(raw_equipment_and_streams)
        else:
            equipment_and_streams_markdown = ""

        sections = [
            ("Problem Statement", final_state.get("problem_statement", "")),
            ("Process Requirements", final_state.get("process_requirements", "")),
            ("Concept Detail", final_state.get("selected_concept_details", "")),
            ("Design Basis", final_state.get("design_basis", "")),
            ("Flowsheet Description", final_state.get("flowsheet_description", "")),
            ("Equipment and Streams List", equipment_and_streams_markdown),
            ("Safety & Risk Assessment", final_state.get("safety_risk_analyst_report", "")),
            ("Project Manager Report", final_state.get("project_manager_report", "")),
        ]
        return sections

    def _write_markdown_report(self, final_state: Dict[str, Any], filename: str) -> None:
        sections = self._compose_report_sections(final_state)

        output_lines: list[str] = []
        for title, content in sections:
            if not content:
                continue
            output_lines.append(f"# {title}")
            output_lines.append(content.strip())
            output_lines.append("")  # Blank line separator

        report_text = "\n".join(output_lines).rstrip() + "\n"

        output_path = Path(filename)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(report_text, encoding="utf-8")

    def _write_word_report(self, final_state: Dict[str, Any], filename: str) -> None:
        if Document is None:
            raise ImportError(
                "python-docx is required to export Word reports. Install with `pip install python-docx`."
            )

        sections = self._compose_report_sections(final_state)
        document = Document()

        output_text = ""
        for title, content in sections:
            if not content:
                continue
            # document.add_heading(title, level=1)
            # document.add_paragraph(content.strip())
            output_text += f"# {title}\n{content.strip()}\n\n"

        output_path = Path(filename)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        # document.save(output_path)
        self._export_markdown_to_word(output_text, output_path)

    def _export_markdown_to_word(self, markdown_string: str, output_filename: str = "report.docx"):
        """
        Converts a Markdown string into a formatted Word (.docx) document.

        Args:
            markdown_string: The string containing your final markdown text.
            output_filename: The name of the Word file to create (e.g., "my_document.docx").
        """
        try:
            # Use pypandoc to convert the string
            # 'md' is the source format (Markdown)
            # 'docx' is the target format (Word)
            pypandoc.convert_text(
                markdown_string,
                'docx',
                format='md',
                outputfile=output_filename,
                extra_args=[
                    f"--reference-doc={self.config.get('save_dir')}/template.docx",
                ]
            )
            
            print(f"\nSuccessfully exported Word document to: {os.path.abspath(output_filename)}\n")

        except FileNotFoundError:
            print("\n--- ERROR ---")
            print("Pandoc executable not found.")
            print("Please ensure Pandoc is installed on your system and available in your PATH.")
        except Exception as e:
            print(f"\nAn error occurred during conversion: {e}")
